from docx import Document
import re
import sys

inputfile = str(sys.argv[1]) # don't give extension!
document = Document()
#document.add_heading(i, 0)
myfile = open(inputfile+'.txt').read()
p = document.add_paragraph(myfile)
outputfile = inputfile + '.docx'
document.save(outputfile)
